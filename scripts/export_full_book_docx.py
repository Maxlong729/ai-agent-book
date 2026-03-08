#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""统一导出全书 Word。

流程：
1. 合并分章稿 -> full-book.md
2. 修正弯引号
3. 准备 reference.docx（修复 Heading 4 斜体、Caption 样式）
4. 调用 pandoc 导出 docx
5. 后处理 docx：图片居中、图片标题置于下方并居中、标题后空一行
"""

from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph

import fix_quotes
import merge_full_book


ROOT = Path(__file__).resolve().parent.parent
MANUSCRIPT_DIR = ROOT / "manuscript"
FULL_BOOK_MD = MANUSCRIPT_DIR / "full-book.md"
FULL_BOOK_DOCX = MANUSCRIPT_DIR / "full-book.docx"
REFERENCE_DOCX = ROOT / "templates" / "reference.docx"


def ensure_reference_docx() -> None:
    REFERENCE_DOCX.parent.mkdir(parents=True, exist_ok=True)
    if not REFERENCE_DOCX.exists():
        result = subprocess.run(
            ["pandoc", "--print-default-data-file", "reference.docx"],
            check=True,
            capture_output=True,
        )
        REFERENCE_DOCX.write_bytes(result.stdout)

    doc = Document(str(REFERENCE_DOCX))

    heading4 = doc.styles["Heading 4"]
    heading4.font.italic = False
    heading4.paragraph_format.space_before = Pt(12)

    heading5 = doc.styles["Heading 5"]
    heading5.font.italic = False
    heading5.paragraph_format.space_before = Pt(10)

    if "Heading 3" in [style.name for style in doc.styles]:
        heading3 = doc.styles["Heading 3"]
        heading3.font.italic = False
        heading3.paragraph_format.space_before = Pt(14)

    if "Heading 6" in [style.name for style in doc.styles]:
        heading6 = doc.styles["Heading 6"]
        heading6.font.italic = False
        heading6.paragraph_format.space_before = Pt(8)

    caption = next((style for style in doc.styles if style.name == "Caption"), None)
    if caption is None:
        caption = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)

    caption.font.italic = False
    caption.font.size = Pt(10.5)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(0)

    doc.save(str(REFERENCE_DOCX))


def collect_existing_image_captions(md_text: str) -> list[str]:
    captions: list[str] = []
    pattern = re.compile(r"^!\[(.*?)\]\((.*?)\)\s*$")
    for line in md_text.splitlines():
        match = pattern.match(line.strip())
        if not match:
            continue
        caption, rel_path = match.groups()
        image_path = MANUSCRIPT_DIR / rel_path
        if not image_path.exists():
            continue
        clean_caption = caption.strip() or image_path.stem.replace("_", " ")
        captions.append(clean_caption)
    return captions


def paragraph_has_image(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


def postprocess_docx(docx_path: Path, captions: list[str]) -> None:
    doc = Document(str(docx_path))

    heading4 = doc.styles["Heading 4"]
    heading4.font.italic = False
    heading4.paragraph_format.space_before = Pt(12)
    heading5 = doc.styles["Heading 5"]
    heading5.font.italic = False
    heading5.paragraph_format.space_before = Pt(10)
    if "Heading 3" in [style.name for style in doc.styles]:
        heading3 = doc.styles["Heading 3"]
        heading3.font.italic = False
        heading3.paragraph_format.space_before = Pt(14)
    if "Heading 6" in [style.name for style in doc.styles]:
        heading6 = doc.styles["Heading 6"]
        heading6.font.italic = False
        heading6.paragraph_format.space_before = Pt(8)
    caption_style = next((style for style in doc.styles if style.name == "Caption"), None)

    image_paragraphs = [p for p in doc.paragraphs if paragraph_has_image(p)]
    for paragraph, caption_text in zip(image_paragraphs, captions):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        caption_para = insert_paragraph_after(paragraph, caption_text, style=caption_style)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.paragraph_format.space_before = Pt(0)
        caption_para.paragraph_format.space_after = Pt(0)

        spacer_para = insert_paragraph_after(caption_para, "")
        spacer_para.paragraph_format.space_before = Pt(0)
        spacer_para.paragraph_format.space_after = Pt(0)

        # 删除 pandoc 已生成的重复图片标题，只保留我们统一插入的居中标题。
        next_para = spacer_para._p.getnext()
        if next_para is not None:
            duplicate = Paragraph(next_para, spacer_para._parent)
            style_name = duplicate.style.name if duplicate.style is not None else ""
            if duplicate.text.strip() == caption_text.strip() or style_name in {"Image Caption", "Caption"}:
                delete_paragraph(duplicate)

    # 让表格尽量整体留在同一页：表格前段落与表格保持一起，表格行不拆页。
    for paragraph in doc.paragraphs:
        paragraph_format = paragraph.paragraph_format
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            paragraph_format.keep_with_next = True
            paragraph_format.keep_together = True
            paragraph_format.space_before = Pt(12)

    for table in doc.tables:
        tbl = table._tbl
        for row in table.rows:
            trPr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            trPr.append(cant_split)
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_with_next = True
                    p.paragraph_format.keep_together = True

    doc.save(str(docx_path))


def export_docx() -> None:
    subprocess.run(
        [
            "pandoc",
            str(FULL_BOOK_MD),
            "-o",
            str(FULL_BOOK_DOCX),
            "--from",
            "markdown",
            "--to",
            "docx",
            "--resource-path",
            str(MANUSCRIPT_DIR),
            "--reference-doc",
            str(REFERENCE_DOCX),
        ],
        cwd=str(ROOT),
        check=True,
    )


def main() -> None:
    merge_full_book.main()

    text = FULL_BOOK_MD.read_text(encoding="utf-8")
    FULL_BOOK_MD.write_text(fix_quotes.fix_quotes(text), encoding="utf-8")

    ensure_reference_docx()
    captions = collect_existing_image_captions(FULL_BOOK_MD.read_text(encoding="utf-8"))
    export_docx()
    postprocess_docx(FULL_BOOK_DOCX, captions)

    print(f"已生成：{FULL_BOOK_DOCX}")


if __name__ == "__main__":
    main()
