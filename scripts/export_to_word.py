#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 Markdown 文件转换为 Word 文档，保持格式一致并嵌入图片
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_chinese_font(run, font_name='微软雅黑', font_size=12):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


def create_document_styles(doc):
    """创建文档样式"""
    styles = doc.styles
    
    # 正文样式
    style = styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)


def parse_markdown_line(line, doc, images_dir, in_table=False, in_code_block=False):
    """解析单行 Markdown 并添加到文档"""
    stripped = line.strip()
    
    # 空行
    if not stripped:
        return None, in_table, in_code_block
    
    # 代码块标记
    if stripped.startswith('```'):
        return 'code_block_toggle', in_table, not in_code_block
    
    # 如果在代码块内
    if in_code_block:
        return ('code', line.rstrip()), in_table, in_code_block
    
    # 标题
    if stripped.startswith('#'):
        match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if match:
            level = len(match.group(1))
            text = match.group(2)
            return ('heading', level, text), False, in_code_block
    
    # 图片
    img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
    if img_match:
        alt_text = img_match.group(1)
        img_path = img_match.group(2)
        return ('image', alt_text, img_path), in_table, in_code_block
    
    # 表格行
    if stripped.startswith('|') and stripped.endswith('|'):
        # 检查是否是分隔行
        if re.match(r'^\|[\s\-:|]+\|$', stripped):
            return ('table_separator',), True, in_code_block
        cells = [cell.strip() for cell in stripped.split('|')[1:-1]]
        return ('table_row', cells), True, in_code_block
    
    # 列表项
    if re.match(r'^[\*\-]\s+', stripped):
        text = re.sub(r'^[\*\-]\s+', '', stripped)
        return ('bullet', text), False, in_code_block
    
    # 有序列表
    if re.match(r'^\d+\.\s+', stripped):
        text = re.sub(r'^\d+\.\s+', '', stripped)
        return ('numbered', text), False, in_code_block
    
    # 引用
    if stripped.startswith('>'):
        text = stripped[1:].strip()
        return ('quote', text), in_table, in_code_block
    
    # 水平线
    if re.match(r'^[-*_]{3,}$', stripped):
        return ('hr',), in_table, in_code_block
    
    # 普通段落
    return ('paragraph', stripped), False, in_code_block


def add_formatted_text(paragraph, text):
    """添加带格式的文本（处理加粗、斜体、行内代码等）"""
    # 处理加粗、斜体、代码等
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\$\$[^$]+\$\$|\$[^$]+\$)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            set_chinese_font(run, font_size=11)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            set_chinese_font(run, font_size=11)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif part.startswith('$$') and part.endswith('$$'):
            # 块级公式，简化处理
            run = paragraph.add_run(part[2:-2])
            run.italic = True
            set_chinese_font(run, font_size=11)
        elif part.startswith('$') and part.endswith('$'):
            # 行内公式，简化处理
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            set_chinese_font(run, font_size=11)
        else:
            run = paragraph.add_run(part)
            set_chinese_font(run, font_size=11)


def process_markdown_file(md_path, doc, images_dir):
    """处理单个 Markdown 文件"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    in_table = False
    in_code_block = False
    table_data = []
    code_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        result, in_table_new, in_code_block_new = parse_markdown_line(
            line, doc, images_dir, in_table, in_code_block
        )
        
        # 代码块切换
        if result == 'code_block_toggle':
            if in_code_block and code_lines:
                # 结束代码块，添加到文档
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                for code_line in code_lines:
                    run = p.add_run(code_line + '\n')
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                code_lines = []
            in_code_block = in_code_block_new
            i += 1
            continue
        
        # 代码行
        if result and result[0] == 'code':
            code_lines.append(result[1])
            in_code_block = in_code_block_new
            i += 1
            continue
        
        # 表格处理
        if in_table and not in_table_new and table_data:
            # 表格结束，创建表格
            if len(table_data) > 0:
                num_cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.style = 'Table Grid'
                for row_idx, row_data in enumerate(table_data):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = table.cell(row_idx, col_idx)
                            cell.text = ''
                            p = cell.paragraphs[0]
                            add_formatted_text(p, cell_text)
                doc.add_paragraph()  # 表格后空行
            table_data = []
        
        in_table = in_table_new
        in_code_block = in_code_block_new
        
        if result is None:
            doc.add_paragraph()
            i += 1
            continue
        
        if result[0] == 'heading':
            level, text = result[1], result[2]
            if level == 1:
                p = doc.add_heading(text, level=0)
            else:
                p = doc.add_heading(text, level=min(level, 4))
            # 设置标题字体
            for run in p.runs:
                set_chinese_font(run, font_size=16 - level * 2 if level < 4 else 11)
        
        elif result[0] == 'image':
            alt_text, img_path = result[1], result[2]
            # 构建完整图片路径
            full_img_path = os.path.join(images_dir, os.path.basename(img_path))
            if not os.path.exists(full_img_path):
                full_img_path = os.path.join(os.path.dirname(md_path), img_path)
            
            if os.path.exists(full_img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(full_img_path, width=Inches(5.5))
                if alt_text:
                    cap = doc.add_paragraph(alt_text)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.runs[0].italic = True
                    set_chinese_font(cap.runs[0], font_size=9)
            else:
                p = doc.add_paragraph(f'[图片: {img_path}]')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif result[0] == 'table_row':
            table_data.append(result[1])
        
        elif result[0] == 'table_separator':
            pass  # 忽略分隔行
        
        elif result[0] == 'bullet':
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, result[1])
        
        elif result[0] == 'numbered':
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, result[1])
        
        elif result[0] == 'quote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(0)
            add_formatted_text(p, result[1])
            for run in p.runs:
                run.italic = True
        
        elif result[0] == 'hr':
            p = doc.add_paragraph('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif result[0] == 'paragraph':
            p = doc.add_paragraph()
            add_formatted_text(p, result[1])
        
        i += 1
    
    # 处理文件末尾的表格
    if table_data:
        num_cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=num_cols)
        table.style = 'Table Grid'
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.cell(row_idx, col_idx)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    add_formatted_text(p, cell_text)


def main():
    """主函数"""
    base_dir = Path('/Users/max/cursor-project/ai-agent-book/manuscript')
    images_dir = base_dir / 'images'
    output_path = base_dir / 'AI智能体工作流_前言及前四章.docx'
    
    # 要处理的文件列表
    md_files = [
        '00-前言.md',
        '01-第01章.md',
        '02-第02章.md',
        '03-第03章.md',
        '04-第04章.md',
    ]
    
    # 创建文档
    doc = Document()
    create_document_styles(doc)
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # 处理每个文件
    for idx, md_file in enumerate(md_files):
        md_path = base_dir / md_file
        if md_path.exists():
            print(f'处理: {md_file}')
            process_markdown_file(str(md_path), doc, str(images_dir))
            
            # 在章节之间添加分页符（除了最后一章）
            if idx < len(md_files) - 1:
                doc.add_page_break()
        else:
            print(f'文件不存在: {md_file}')
    
    # 保存文档
    doc.save(str(output_path))
    print(f'\n导出完成: {output_path}')


if __name__ == '__main__':
    main()
